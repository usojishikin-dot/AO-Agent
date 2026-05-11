"""
build_docs.py
Converts API_DOCUMENTATION.md into a properly formatted Word document.
Run: python build_docs.py
"""

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = Path("API_DOCUMENTATION.md")
OUT_FILE = Path("API_DOCUMENTATION.docx")

# ── Colour palette ────────────────────────────────────────────────
DARK_BG      = RGBColor(0x1E, 0x1E, 0x2E)   # code block background
CODE_FG      = RGBColor(0xD4, 0xD4, 0xD4)   # code text
ACCENT       = RGBColor(0x00, 0x70, 0xC0)   # headings / links
TABLE_HEADER = RGBColor(0x00, 0x46, 0x8C)   # table header bg
TABLE_ALT    = RGBColor(0xE9, 0xF1, 0xFA)   # alternating row
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)
MUTED        = RGBColor(0x66, 0x66, 0x66)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_para_border_bottom(para, color="CCCCCC", size="6"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]]):
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, TABLE_HEADER)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = TABLE_ALT if row_idx % 2 == 0 else WHITE
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            # Render inline code spans inside table cells
            add_inline_text(para, cell_text, base_size=9)

    doc.add_paragraph()  # spacing after table


def add_inline_text(para, text: str, base_size: int = 10, bold: bool = False):
    """Render a string that may contain `backtick code` spans into a paragraph."""
    segments = re.split(r'(`[^`]+`)', text)
    for seg in segments:
        if seg.startswith("`") and seg.endswith("`"):
            inner = seg[1:-1]
            run = para.add_run(inner)
            run.font.name = "Courier New"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4C)
        else:
            if seg:
                run = para.add_run(seg)
                run.bold = bold
                run.font.size = Pt(base_size)


def add_code_block(doc: Document, code: str, lang: str = ""):
    """Add a shaded code block paragraph."""
    # Remove fenced markers
    lines = code.strip().split("\n")
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # Grey shading via XML
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F4F4F4")
    pPr.append(shd)

    run = para.add_run("\n".join(lines))
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    doc.add_paragraph()


def parse_table(lines: list[str]):
    """Parse a markdown table into (headers, rows)."""
    headers = [c.strip().strip("|").strip() for c in lines[0].split("|") if c.strip()]
    rows = []
    for line in lines[2:]:  # skip separator
        if not line.strip():
            continue
        cells = [c.strip() for c in line.split("|") if c != ""]
        # pad / trim to header length
        while len(cells) < len(headers):
            cells.append("")
        rows.append(cells[:len(headers)])
    return headers, rows


def build_document(md_text: str) -> Document:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Define named styles ───────────────────────────────────────
    styles = doc.styles

    def ensure_style(name, base="Normal"):
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, 1)

    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Skip horizontal rules ─────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()):
            para = doc.add_paragraph()
            set_para_border_bottom(para)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text  = heading_match.group(2).strip()
            # Strip markdown bold/italic
            text  = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text  = re.sub(r"\*(.+?)\*",     r"\1", text)

            if level == 1:
                para = doc.add_heading(text, level=0)
                para.runs[0].font.color.rgb = ACCENT
                para.runs[0].font.size = Pt(22)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif level == 2:
                para = doc.add_heading(text, level=1)
                para.runs[0].font.color.rgb = ACCENT
                para.runs[0].font.size = Pt(16)
            elif level == 3:
                para = doc.add_heading(text, level=2)
                para.runs[0].font.color.rgb = RGBColor(0x20, 0x60, 0xA0)
                para.runs[0].font.size = Pt(13)
            else:
                para = doc.add_heading(text, level=3)
                para.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                para.runs[0].font.size = Pt(11)
            i += 1
            continue

        # ── Fenced code blocks ────────────────────────────────────
        fence_match = re.match(r"^```(\w*)", line)
        if fence_match:
            lang = fence_match.group(1)
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, "\n".join(code_lines), lang)
            continue

        # ── Alert blockquotes (> [!NOTE] etc.) ───────────────────
        if line.startswith("> [!"):
            alert_type = re.match(r"> \[!(\w+)\]", line)
            label = alert_type.group(1) if alert_type else "NOTE"
            i += 1
            content_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                content_lines.append(lines[i].lstrip("> ").strip())
                i += 1
            full_text = " ".join(content_lines)
            icon = {"NOTE": "ℹ", "IMPORTANT": "❗", "WARNING": "⚠", "CAUTION": "🚨", "TIP": "💡"}.get(label, "•")
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            shd  = OxmlElement("w:shd")
            fill = {"NOTE": "E8F4FD", "IMPORTANT": "FFF3CD", "WARNING": "FFF3CD", "CAUTION": "FDECEA", "TIP": "E8F8F0"}.get(label, "F0F0F0")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            pPr.append(shd)
            para.paragraph_format.left_indent = Cm(0.5)
            run = para.add_run(f"{icon} {label}:  ")
            run.bold = True
            run.font.size = Pt(9.5)
            add_inline_text(para, full_text, base_size=9)
            doc.add_paragraph()
            continue

        # ── Markdown tables ───────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                headers, rows = parse_table(table_lines)
                add_styled_table(doc, headers, rows)
            continue

        # ── Bullet lists (- or *) ─────────────────────────────────
        bullet_match = re.match(r"^(\s*)[-*]\s+(.+)", line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text   = bullet_match.group(2)
            style  = "List Bullet 2" if indent > 0 else "List Bullet"
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_paragraph(style="List Bullet")
            add_inline_text(para, text, base_size=10)
            i += 1
            continue

        # ── Numbered lists ────────────────────────────────────────
        num_match = re.match(r"^\d+\.\s+(.+)", line)
        if num_match:
            text = num_match.group(1)
            try:
                para = doc.add_paragraph(style="List Number")
            except KeyError:
                para = doc.add_paragraph()
            add_inline_text(para, text, base_size=10)
            i += 1
            continue

        # ── Empty lines ───────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────
        # Skip table-of-contents anchor links
        if re.match(r"^\d+\. \[", line):
            i += 1
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        add_inline_text(para, line, base_size=10)
        i += 1

    return doc


if __name__ == "__main__":
    if not MD_FILE.exists():
        print(f"ERROR: '{MD_FILE}' not found. Run from the project root.")
        raise SystemExit(1)

    md_text = MD_FILE.read_text(encoding="utf-8")
    print("Parsing markdown...")
    document = build_document(md_text)
    document.save(OUT_FILE)
    print(f"SUCCESS: Word document saved -> {OUT_FILE.resolve()}")
